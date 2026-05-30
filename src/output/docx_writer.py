"""Write transcribed segments to a DOCX file via python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from transcription.engine import TranscribedSegment


def _format_time(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = seconds % 60
    return f"{h:02d}:{m:02d}:{s:06.3f}"


def write(
    segments: list["TranscribedSegment"],
    path: Path | str,
    fields: dict[str, bool] | None = None,
) -> Path:
    """Write *segments* to *path* as a DOCX document.

    Each segment occupies one or more paragraphs (one per enabled field).
    A blank paragraph separates segments.

    Parameters
    ----------
    segments:
        List of TranscribedSegment objects.
    path:
        Output .docx path.
    fields:
        Dict of ``{field_name: bool}`` controlling which fields appear.
    """
    from docx import Document

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    _fields = fields or {
        "timestamp": True,
        "speaker": True,
        "language": True,
        "confidence": True,
        "text": True,
        "translation": False,
    }

    doc = Document()

    for i, seg in enumerate(segments):
        if _fields.get("timestamp", True):
            doc.add_paragraph(f"{_format_time(seg.start)} --> {_format_time(seg.end)}")
        if _fields.get("speaker", True):
            doc.add_paragraph(f"Speaker: {seg.speaker_id}")
        if _fields.get("language", True):
            doc.add_paragraph(f"Language: {seg.language}")
        if _fields.get("confidence", True):
            doc.add_paragraph(f"Confidence: {seg.confidence:.2f}")
        if _fields.get("text", True):
            doc.add_paragraph(f"Text: {seg.text}")
        if _fields.get("translation", False) and seg.translated_text:
            doc.add_paragraph(f"Translation: {seg.translated_text}")
        if i < len(segments) - 1:
            doc.add_paragraph("")  # blank separator

    doc.save(str(path))
    return path
