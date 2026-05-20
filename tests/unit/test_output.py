"""Unit tests for output writers and naming — CHK-43 to 50."""

import json
from pathlib import Path

import pytest

from config.store import ConfigStore
from output import naming, txt_writer, srt_writer, json_writer, clipboard_writer, docx_writer
from output.clipboard_writer import ClipboardNotAvailableError
from transcription.engine import TranscribedSegment


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def make_seg(text: str, start: float = 0.0, end: float = 3.0,
             speaker: str = "Speaker 1", language: str = "English",
             language_code: str = "en", confidence: float = 0.9,
             no_speech_prob: float = 0.1, bad_audio: bool = False) -> TranscribedSegment:
    return TranscribedSegment(
        speaker_id=speaker,
        start=start, end=end,
        text=text, language=language, language_code=language_code,
        confidence=confidence, no_speech_prob=no_speech_prob,
        bad_audio=bad_audio,
    )


def three_segs():
    return [
        make_seg("Hello world", start=0.0, end=3.0),
        make_seg("How are you", start=3.5, end=6.0, speaker="Speaker 2"),
        make_seg("I am fine", start=6.5, end=9.0),
    ]


# ---------------------------------------------------------------------------
# CHK-43: TXT writer — 3 segments, all 6 fields, correct separators
# ---------------------------------------------------------------------------

def test_txt_writer_creates_file(tmp_path):
    path = tmp_path / "out.txt"
    txt_writer.write(three_segs(), path)
    assert path.exists()


def test_txt_writer_contains_all_segments(tmp_path):
    path = tmp_path / "out.txt"
    txt_writer.write(three_segs(), path)
    content = path.read_text(encoding="utf-8")
    assert "Hello world" in content
    assert "How are you" in content
    assert "I am fine" in content


def test_txt_writer_blank_line_separators(tmp_path):
    path = tmp_path / "out.txt"
    txt_writer.write(three_segs(), path)
    content = path.read_text(encoding="utf-8")
    assert "\n\n" in content


def test_txt_writer_speaker_field(tmp_path):
    path = tmp_path / "out.txt"
    txt_writer.write(three_segs(), path)
    content = path.read_text(encoding="utf-8")
    assert "Speaker:" in content


def test_txt_writer_timestamp_field(tmp_path):
    path = tmp_path / "out.txt"
    txt_writer.write(three_segs(), path)
    content = path.read_text(encoding="utf-8")
    assert "-->" in content


# ---------------------------------------------------------------------------
# CHK-44: TXT writer — language field disabled → no language prefix
# ---------------------------------------------------------------------------

def test_txt_writer_no_language_when_disabled(tmp_path):
    path = tmp_path / "out.txt"
    fields = {
        "timestamp": True, "speaker": True,
        "language": False, "confidence": True, "text": True,
    }
    txt_writer.write(three_segs(), path, fields=fields)
    content = path.read_text(encoding="utf-8")
    assert "Language:" not in content


def test_txt_writer_language_present_when_enabled(tmp_path):
    path = tmp_path / "out.txt"
    txt_writer.write(three_segs(), path)
    content = path.read_text(encoding="utf-8")
    assert "Language: English" in content


# ---------------------------------------------------------------------------
# CHK-45: SRT writer — valid format
# ---------------------------------------------------------------------------

def test_srt_writer_creates_file(tmp_path):
    path = tmp_path / "out.srt"
    srt_writer.write(three_segs(), path)
    assert path.exists()


def test_srt_writer_sequential_numbers(tmp_path):
    path = tmp_path / "out.srt"
    srt_writer.write(three_segs(), path)
    content = path.read_text(encoding="utf-8")
    assert content.startswith("1\n")
    assert "2\n" in content
    assert "3\n" in content


def test_srt_writer_timestamp_format(tmp_path):
    path = tmp_path / "out.srt"
    srt_writer.write(three_segs(), path)
    content = path.read_text(encoding="utf-8")
    # SRT timestamp: HH:MM:SS,mmm --> HH:MM:SS,mmm
    assert "00:00:00,000 --> 00:00:03,000" in content


def test_srt_writer_speaker_in_text(tmp_path):
    path = tmp_path / "out.srt"
    srt_writer.write(three_segs(), path)
    content = path.read_text(encoding="utf-8")
    assert "Speaker 1: Hello world" in content


def test_srt_writer_no_language_prefix(tmp_path):
    path = tmp_path / "out.srt"
    seg = make_seg("Hello", language="Finnish", language_code="fi")
    srt_writer.write([seg], path)
    content = path.read_text(encoding="utf-8")
    assert "Finnish" not in content
    assert "[fi]" not in content


# ---------------------------------------------------------------------------
# CHK-46: JSON writer — all 6 fields present
# ---------------------------------------------------------------------------

def test_json_writer_creates_valid_json(tmp_path):
    path = tmp_path / "out.json"
    json_writer.write(three_segs(), path)
    data = json.loads(path.read_text(encoding="utf-8"))
    assert isinstance(data, list)
    assert len(data) == 3


def test_json_writer_all_fields_present(tmp_path):
    path = tmp_path / "out.json"
    json_writer.write(three_segs(), path)
    data = json.loads(path.read_text(encoding="utf-8"))
    required = {"start", "end", "speaker", "language", "confidence", "text",
                "no_speech_prob", "bad_audio", "low_confidence", "language_code"}
    for item in data:
        assert required.issubset(item.keys())


def test_json_writer_field_values(tmp_path):
    path = tmp_path / "out.json"
    json_writer.write([make_seg("Hello", start=1.0, end=4.0)], path)
    data = json.loads(path.read_text(encoding="utf-8"))
    assert data[0]["start"] == 1.0
    assert data[0]["end"] == 4.0
    assert data[0]["text"] == "Hello"
    assert data[0]["speaker"] == "Speaker 1"


# ---------------------------------------------------------------------------
# CHK-47: DOCX writer — opens without error, segment count matches
# ---------------------------------------------------------------------------

def test_docx_writer_creates_file(tmp_path):
    path = tmp_path / "out.docx"
    docx_writer.write(three_segs(), path)
    assert path.exists()


def test_docx_writer_readable_by_python_docx(tmp_path):
    from docx import Document
    path = tmp_path / "out.docx"
    docx_writer.write(three_segs(), path)
    doc = Document(str(path))
    # Should not raise; doc must have some paragraphs
    assert len(doc.paragraphs) > 0


def test_docx_writer_segment_text_present(tmp_path):
    from docx import Document
    path = tmp_path / "out.docx"
    docx_writer.write(three_segs(), path)
    doc = Document(str(path))
    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "Hello world" in all_text
    assert "How are you" in all_text


# ---------------------------------------------------------------------------
# CHK-48: naming — _WSP convention and collision avoidance
# ---------------------------------------------------------------------------

def test_naming_basic(tmp_path):
    path = naming.make_output_path(tmp_path / "interview.mp3", ".txt", tmp_path)
    assert path.name == "interview_WSP.txt"


def test_naming_collision_second_run(tmp_path):
    (tmp_path / "interview_WSP.txt").write_text("")
    path = naming.make_output_path(tmp_path / "interview.mp3", ".txt", tmp_path)
    assert path.name == "interview_WSP_2.txt"


def test_naming_collision_third_run(tmp_path):
    (tmp_path / "interview_WSP.txt").write_text("")
    (tmp_path / "interview_WSP_2.txt").write_text("")
    path = naming.make_output_path(tmp_path / "interview.mp3", ".txt", tmp_path)
    assert path.name == "interview_WSP_3.txt"


# ---------------------------------------------------------------------------
# CHK-49: naming — split streams _part1, _part2
# ---------------------------------------------------------------------------

def test_naming_split_part1(tmp_path):
    path = naming.make_output_path(tmp_path / "interview.mp3", ".txt", tmp_path, part=1)
    assert path.name == "interview_WSP_part1.txt"


def test_naming_split_part2(tmp_path):
    path = naming.make_output_path(tmp_path / "interview.mp3", ".txt", tmp_path, part=2)
    assert path.name == "interview_WSP_part2.txt"


def test_naming_split_part_collision(tmp_path):
    (tmp_path / "interview_WSP_part1.txt").write_text("")
    path = naming.make_output_path(tmp_path / "interview.mp3", ".txt", tmp_path, part=1)
    assert path.name == "interview_WSP_part1_2.txt"


# ---------------------------------------------------------------------------
# CHK-50: ClipboardNotAvailableError for source_type="file"
# ---------------------------------------------------------------------------

def test_clipboard_raises_for_file_source():
    segs = [make_seg("hello")]
    with pytest.raises(ClipboardNotAvailableError, match="file"):
        clipboard_writer.write(segs, source_type="file")


def test_clipboard_raises_for_batch_source():
    segs = [make_seg("hello")]
    with pytest.raises(ClipboardNotAvailableError):
        clipboard_writer.write(segs, source_type="batch")


def test_clipboard_error_message_mentions_spec():
    segs = [make_seg("hello")]
    with pytest.raises(ClipboardNotAvailableError, match="8.1.c"):
        clipboard_writer.write(segs, source_type="file")
